"""
resume_parser.py
Extracts raw text from PDF or DOCX resume files.
"""

import os
from PyPDF2 import PdfReader
import docx


def extract_text_from_pdf(filepath):
    """Extract text from a PDF file."""
    text = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {e}")
    return text.strip()


def extract_text_from_docx(filepath):
    """Extract text from a DOCX file."""
    text = ""
    try:
        document = docx.Document(filepath)
        for para in document.paragraphs:
            text += para.text + "\n"
        # also read text inside tables (if any)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {e}")
    return text.strip()


def extract_text(filepath):
    """Detect file type by extension and extract text accordingly."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx",):
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Please upload a PDF, DOCX, or TXT file."
        )
